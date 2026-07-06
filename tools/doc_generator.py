from docx import Document
import os
from datetime import datetime

def generate_doc(sections:dict, review):
    doc = Document()
    doc.add_heading("System Design Report",0)

    for title,content in sections.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    
    doc.add_heading("AI Reflection Review", level=1)
    doc.add_paragraph(review)

    os.makedirs("output",exist_ok=True)
    filename =  datetime.now().strftime("report_%Y%m%d_%H%M%S.docx")

    file_path = os.path.join("output",filename)

    doc.save(file_path)
    return file_path    
